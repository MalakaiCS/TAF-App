"""
bag_filler.py
Generates a bag-filter / media-roll docket as a Word document (.docx),
then tries to export it to PDF via PowerShell Word COM.

Data model for a single bag line-item (dict keys):
  product_type  : str   – one of BAG_PRODUCT_TYPES
  quantity      : int
  media         : str   – from BAG_MEDIA_TYPES or ROLL_MEDIA_TYPES
  width         : int   – mm  (0 = N/A for rolls)
  height        : int   – mm
  depth         : int   – mm  (0 = N/A for frames / rolls)
  on_wire       : bool
  gelled        : bool
  special_size  : bool
  label_suffix  : str   – "LONG" | "SHORT" | ""
  efficiency    : str   – MPHE only: "60-65%" etc.
  roll_width    : str   – e.g. "2100mm"   (Media Roll only)
  roll_length   : str   – e.g. "20m"      (Media Roll only)
  description   : str   – free-text for Other / Cut Pads
  part_number   : str   – auto-generated (may be overridden by user)
  notes         : str
"""
from __future__ import annotations
import os
import re
import subprocess
import platform
import atexit
import threading as _threading
from pathlib import Path
from typing import List, Dict, Any, Optional

# ── Persistent Word COM cache ─────────────────────────────────────────────────
# Word.Application is expensive to start (~2-5 s).  We keep one instance alive
# for the whole session and reuse it across bag-docket exports.
_word_app  = None
_word_lock = _threading.Lock()


def _ensure_word_app():
    """Return a live Word.Application COM object, creating one if needed."""
    global _word_app
    with _word_lock:
        if _word_app is not None:
            try:
                _ = _word_app.Version      # quick liveness check
                return _word_app
            except Exception:
                _word_app = None
        try:
            import win32com.client
            app = win32com.client.Dispatch("Word.Application")
            app.Visible       = False
            app.DisplayAlerts = 0          # wdAlertsNone
            _word_app = app
        except Exception:
            _word_app = None
        return _word_app


def _quit_word():
    """Called at process exit — gracefully close the cached Word instance."""
    global _word_app
    with _word_lock:
        if _word_app is not None:
            try:
                _word_app.Quit(0)          # 0 = wdDoNotSaveChanges
            except Exception:
                pass
            _word_app = None

atexit.register(_quit_word)


def _export_word_pdf_via_cached_word(docx_path: str, pdf_path: str) -> bool:
    """Export docx→PDF using the persistent Word COM object.  Returns True on success."""
    for attempt in range(2):
        app = _ensure_word_app()
        if app is None:
            return False
        try:
            doc = app.Documents.Open(docx_path)
            doc.SaveAs2(pdf_path, 17)      # 17 = wdFormatPDF
            doc.Close(False)
            return os.path.exists(pdf_path)
        except Exception:
            with _word_lock:
                _word_app = None           # drop stale reference and retry once
            if attempt == 1:
                return False
    return False


def prewarm_word():
    """Call on a background thread at app startup to hide Word's first-launch cost."""
    _ensure_word_app()

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# ── Catalogue constants ───────────────────────────────────────────────────────

BAG_PRODUCT_TYPES = [
    "3-Peak",
    "2-Wedge",
    "4-Point",
    "MPHE 8-Pocket",
    "MPHE 4-Pocket",
    "HEPA",
    "Mounting Frame",
    "Media Roll",
    "Cut Pads",
    "Other",
]

BAG_MEDIA_TYPES  = ["G4", "F5", "F6", "F7", "F8"]
ROLL_MEDIA_TYPES = ["G4", "180", "GREY", "Paintstop"]
ROLL_WIDTHS      = ["600mm", "1000mm", "2100mm"]
ROLL_LENGTHS     = ["8m", "20m", "40m"]
MPHE_EFFICIENCIES = ["", "60-65%", "80-85%", "90-95%"]

# product types that are "bag" filters (need W/H/D + media)
_BAG_TYPES  = {"3-Peak", "2-Wedge", "4-Point", "MPHE 8-Pocket", "MPHE 4-Pocket", "HEPA"}
# types that use bag media dropdown
_MEDIA_TYPES_BAG  = {"3-Peak", "2-Wedge", "4-Point", "HEPA"}
_MEDIA_TYPES_MPHE = {"MPHE 8-Pocket", "MPHE 4-Pocket"}

# Standard size presets: {product_type: [(label, width, height, depth|None), ...]}
STANDARD_SIZES: Dict[str, list] = {
    "3-Peak": [
        ("Full  610×610×570  (Long)",    610, 610, 570),
        ("Full  610×610×340  (Short)",   610, 610, 340),
        ("Full  610×610×600  (X-Long)",  610, 610, 600),
        ("Half  305×610×570  (Long)",    305, 610, 570),
        ("Half  305×610×340  (Short)",   305, 610, 340),
        ("Custom",                       None, None, None),
    ],
    "2-Wedge": [
        ("Full  610×610×340",            610, 610, 340),
        ("Full  610×610×570  (Long)",    610, 610, 570),
        ("Full  610×610×600  (X-Long)",  610, 610, 600),
        ("Half  305×610×340",            305, 610, 340),
        ("Custom",                       None, None, None),
    ],
    "4-Point": [
        ("Full  610×610×340",            610, 610, 340),
        ("Full  610×610×600  (Long)",    610, 610, 600),
        ("Half  305×610×340",            305, 610, 340),
        ("Half  305×610×600  (Long)",    305, 610, 600),
        ("Custom",                       None, None, None),
    ],
    "MPHE 8-Pocket": [
        ("Full  595×595×560",            595, 595, 560),
        ("Full  595×595×350  (Short)",   595, 595, 350),
        ("Custom",                       None, None, None),
    ],
    "MPHE 4-Pocket": [
        ("Half  295×595×560",            295, 595, 560),
        ("Half  295×595×350  (Short)",   295, 595, 350),
        ("Custom",                       None, None, None),
    ],
    "HEPA": [
        ("Full  610×610×69",             610, 610,  69),
        ("Full  610×610×292",            610, 610, 292),
        ("Custom",                       None, None, None),
    ],
    "Mounting Frame": [
        ("610×610  Standard",            610, 610, None),
        ("Custom",                       None, None, None),
    ],
}


# ── Part number generation ────────────────────────────────────────────────────

def generate_part_number(item: dict) -> str:
    """Return the auto-generated part number string for a bag line item."""
    pt      = item.get("product_type", "")
    media   = (item.get("media") or "").upper()
    w       = item.get("width")  or 0
    h       = item.get("height") or 0
    d       = item.get("depth")  or 0
    on_wire = bool(item.get("on_wire"))
    gelled  = bool(item.get("gelled"))
    special = bool(item.get("special_size"))

    bag_sfx  = ("W" if on_wire else "") + ("G" if gelled else "")
    spec_sfx = " SPEC" if special else ""

    # helper
    def _d3():
        return f"{int(w)}{int(h)}{int(d)}" if (w and h and d) else ""
    def _d2():
        return f"{int(w)}{int(h)}" if (w and h) else ""

    if pt == "3-Peak":
        dims = _d3()
        return f"3P{media}-{dims}{bag_sfx}{spec_sfx}" if dims else ""

    if pt == "2-Wedge":
        dims = _d3()
        return f"2W{media}-{dims}{bag_sfx}{spec_sfx}" if dims else ""

    if pt == "4-Point":
        dims = _d3()
        return f"4P{media}-{dims}{bag_sfx}{spec_sfx}" if dims else ""

    if pt == "MPHE 8-Pocket":
        dims = _d3()
        return f"MPHE-{media}8P{dims}" if dims else ""

    if pt == "MPHE 4-Pocket":
        dims = _d3()
        return f"MPHE-{media}4P{dims}" if dims else ""

    if pt == "HEPA":
        dims = _d3()
        return f"GSHEP99-{dims}" if dims else ""

    if pt == "Mounting Frame":
        dims = _d2()
        return f"TAFGSMF-{dims}" if dims else ""

    if pt == "Media Roll":
        roll_m = (item.get("media") or "G4").upper()
        rw = str(item.get("roll_width") or "")
        if roll_m == "G4":
            return "TAFTFM390D"
        if roll_m == "180":
            return "TAFUFM180"
        if roll_m == "PAINTSTOP":
            if "600" in rw:  return "TAF-PAINTSTOP 061"
            if "1000" in rw: return "TAF-Paintstop 100"
        return ""

    return ""


# ── Description / display helpers ─────────────────────────────────────────────

def build_label_line(item: dict) -> str:
    """Return the first description line, e.g. '4 x 3-Peak G4 / ON WIRE'."""
    pt     = item.get("product_type", "")
    qty    = item.get("quantity", 1)
    media  = (item.get("media") or "").upper()
    w      = int(item.get("width") or 0)
    on_wire = bool(item.get("on_wire"))
    gelled  = bool(item.get("gelled"))
    lsuffix = (item.get("label_suffix") or "").strip().upper()
    eff     = (item.get("efficiency") or "").strip()

    half_pfx = "½ " if w in (305, 295) else ""

    quals = []
    if lsuffix:  quals.append(lsuffix)
    if on_wire:  quals.append("ON WIRE")
    if gelled:   quals.append("GELLED")
    qual = " / ".join(quals)

    if pt in ("3-Peak", "HEPA"):
        name = f"{half_pfx}{pt} {media}"
    elif pt == "2-Wedge":
        name = f"{half_pfx}2-Wedge {media}"
    elif pt == "4-Point":
        name = f"{half_pfx}4-POINT {media}"
    elif pt in ("MPHE 8-Pocket", "MPHE 4-Pocket"):
        pockets = "8-Pocket" if "8" in pt else "4-Pocket"
        eff_str = f" {eff}" if eff else ""
        name = f"{half_pfx}{pockets} {media}{eff_str}"
    elif pt == "Mounting Frame":
        return f"{qty} x Gal Steel Mounting Frame"
    elif pt == "Media Roll":
        roll_m = (item.get("media") or "G4").upper()
        return f"{qty} x FILTER MEDIA ROLL – {roll_m}"
    elif pt == "Cut Pads":
        desc = item.get("description") or media or ""
        return f"{qty} x CUT PADS – {desc}".rstrip("– ")
    else:
        return f"{qty} x {item.get('description', pt)}"

    line = f"{qty} x {name}"
    if qual:
        line += f" / {qual}"
    return line


def build_dims_line(item: dict) -> str:
    """Return dimension/size line, e.g. '610mm x 610mm x 570mm'."""
    pt = item.get("product_type", "")
    if pt == "Media Roll":
        rw = item.get("roll_width", "2100mm")
        rl = item.get("roll_length", "20m")
        mtrs = rl.replace("m", "")
        return f"{rw} x {mtrs}MTRS Long"

    w = item.get("width")  or 0
    h = item.get("height") or 0
    d = item.get("depth")  or 0

    if pt == "Mounting Frame" or not d:
        return f"{int(w)}mm x {int(h)}mm" if (w and h) else ""

    if pt == "Cut Pads":
        if w and h:
            return f"{int(w)}mm x {int(h)}mm"
        return item.get("description", "")

    return f"{int(w)}mm x {int(h)}mm x {int(d)}mm" if (w and h and d) else ""


def item_summary_short(item: dict) -> str:
    """One-line summary for the treeview."""
    pt = item.get("product_type", "")
    w  = int(item.get("width")  or 0)
    h  = int(item.get("height") or 0)
    d  = int(item.get("depth")  or 0)
    if pt == "Media Roll":
        return f"{item.get('roll_width','')} × {item.get('roll_length','')}"
    if pt == "Mounting Frame":
        return f"{w}×{h}" if (w and h) else ""
    return f"{w}×{h}×{d}" if (w and h and d) else ""


# ── Word document generation ──────────────────────────────────────────────────

_NAVY = (26, 79, 138)   # TAF navy RGB


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _remove_cell_borders(cell):
    """Remove all borders from a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def generate_bag_docket(header: dict, items: List[dict], out_path: str,
                        auto_open: bool = True,
                        item_start: int = 1,
                        grand_total: int = None) -> str:
    """
    Build a Word docket (.docx) for a bag/roll order, save to out_path,
    attempt PDF export via PowerShell Word COM, return the output path.
    Pass auto_open=False to suppress opening the file (for PDF-merge workflows).
    item_start / grand_total set the stock-number range within the full order
    (e.g. item_start=3, grand_total=5 → items numbered 3/5, 4/5, 5/5).
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed.  "
            "Run:  pip install python-docx")

    doc = Document()

    # ── Page setup (A4 portrait) ──────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Styles helper ─────────────────────────────────────────────────────
    def _p(text="", bold=False, size=11, color=None, align=None, space_before=0, space_after=4):
        para = doc.add_paragraph()
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = para.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return para

    def _add_run(para, text, bold=False, size=11, color=None):
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    # ── Header banner table (light grey bg, B&W) ─────────────────────────
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.style = "Table Grid"
    hdr_tbl.autofit = False
    content_width = section.page_width - section.left_margin - section.right_margin
    # left col: 65%, right col: 35%
    left_w  = int(content_width * 0.65)
    right_w = int(content_width * 0.35)
    hdr_tbl.columns[0].width = left_w
    hdr_tbl.columns[1].width = right_w

    lc = hdr_tbl.rows[0].cells[0]
    rc = hdr_tbl.rows[0].cells[1]
    _set_cell_bg(lc, "F0F0F0")
    _set_cell_bg(rc, "F0F0F0")

    # Left: company name + O/N + ATT
    # (cell already has one empty paragraph from table creation — use it)
    lc_para = lc.paragraphs[0]
    lc_para.paragraph_format.space_before = Pt(4)
    lc_para.paragraph_format.space_after  = Pt(2)
    r = lc_para.add_run((header.get("Customer Name") or "").upper())
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 0, 0)

    def _lc_row(label, value):
        p = lc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run_l = p.add_run(f"{label}  ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = RGBColor(60, 60, 60)
        run_v = p.add_run(str(value or ""))
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = RGBColor(0, 0, 0)

    _lc_row("O/N:", header.get("Order Number", ""))
    _lc_row("ATT:", header.get("Attention", ""))
    if header.get("Location"):
        _lc_row("LOCATION:", header.get("Location", ""))
    # small bottom padding
    p_pad = lc.add_paragraph()
    p_pad.paragraph_format.space_after = Pt(4)

    # Right: dates + job  (use the existing first paragraph, then add_paragraph for rest)
    _rc_first = [True]

    def _rc_row(label, value):
        if _rc_first[0]:
            p = rc.paragraphs[0]
            _rc_first[0] = False
        else:
            p = rc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_l = p.add_run(f"{label}  ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = RGBColor(60, 60, 60)
        run_v = p.add_run(str(value or ""))
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = RGBColor(0, 0, 0)

    _rc_row("DATE:", header.get("Date Ordered", ""))
    _rc_row("DUE:", header.get("Date Due", ""))
    if header.get("Job"):
        _rc_row("JOB:", header.get("Job", ""))
    p_pad2 = rc.add_paragraph()
    p_pad2.paragraph_format.space_after = Pt(4)
    p_pad2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Spacer after header
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(6)
    sp.paragraph_format.space_after  = Pt(4)

    # ── Line items (stock-number table style) ──────────────────────────────
    total_shown = grand_total if grand_total is not None else len(items)
    STOCK_W     = int(Cm(1.6))
    DETAIL_W    = int(content_width - STOCK_W)
    _BG_EVEN    = "F5F5F5"
    _BG_ODD     = "FFFFFF"

    for i, item in enumerate(items):
        stock_num = item_start + i
        label = build_label_line(item)
        dims  = build_dims_line(item)
        pn    = (item.get("part_number") or "").strip()
        notes = (item.get("notes") or "").strip()

        # Two-column item table: left = stock badge (navy), right = details
        itbl = doc.add_table(rows=1, cols=2)
        itbl.style   = "Table Grid"
        itbl.autofit = False
        itbl.columns[0].width = STOCK_W
        itbl.columns[1].width = DETAIL_W

        sc = itbl.rows[0].cells[0]
        dc = itbl.rows[0].cells[1]

        # Stock-number cell (grey)
        _set_cell_bg(sc, "DDDDDD")
        _set_cell_valign(sc, "center")
        _set_cell_margins(sc, top=80, bottom=80, left=60, right=60)

        sc_p1 = sc.paragraphs[0]
        sc_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc_p1.paragraph_format.space_before = Pt(0)
        sc_p1.paragraph_format.space_after  = Pt(0)
        r_cur = sc_p1.add_run(str(stock_num))
        r_cur.bold = True; r_cur.font.size = Pt(18)
        r_cur.font.color.rgb = RGBColor(0, 0, 0)

        sc_p2 = sc.add_paragraph()
        sc_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc_p2.paragraph_format.space_before = Pt(0)
        sc_p2.paragraph_format.space_after  = Pt(0)
        r_sl = sc_p2.add_run("/")
        r_sl.font.size = Pt(9)
        r_sl.font.color.rgb = RGBColor(100, 100, 100)

        sc_p3 = sc.add_paragraph()
        sc_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc_p3.paragraph_format.space_before = Pt(0)
        sc_p3.paragraph_format.space_after  = Pt(0)
        r_tot = sc_p3.add_run(str(total_shown))
        r_tot.bold = True; r_tot.font.size = Pt(12)
        r_tot.font.color.rgb = RGBColor(0, 0, 0)

        # Detail cell
        bg = _BG_EVEN if (i % 2 == 0) else _BG_ODD
        _set_cell_bg(dc, bg)
        _set_cell_margins(dc, top=100, bottom=100, left=180, right=120)

        dc_p0 = dc.paragraphs[0]
        dc_p0.paragraph_format.space_before = Pt(2)
        dc_p0.paragraph_format.space_after  = Pt(2)
        rl = dc_p0.add_run(label)
        rl.bold = True; rl.font.size = Pt(12)
        rl.font.color.rgb = RGBColor(0, 0, 0)

        if dims:
            dp = dc.add_paragraph()
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after  = Pt(2)
            rd = dp.add_run(dims)
            rd.font.size = Pt(11)
            rd.font.color.rgb = RGBColor(0, 0, 0)

        if pn:
            pp = dc.add_paragraph()
            pp.paragraph_format.space_before = Pt(0)
            pp.paragraph_format.space_after  = Pt(2)
            r_pl = pp.add_run("P/N:  ")
            r_pl.bold = True; r_pl.font.size = Pt(10)
            r_pl.font.color.rgb = RGBColor(60, 60, 60)
            r_pv = pp.add_run(pn)
            r_pv.font.size = Pt(10)
            r_pv.font.color.rgb = RGBColor(0, 0, 0)

        if notes:
            np_ = dc.add_paragraph()
            np_.paragraph_format.space_before = Pt(0)
            np_.paragraph_format.space_after  = Pt(4)
            r_n = np_.add_run(f"Note:  {notes}")
            r_n.italic = True; r_n.font.size = Pt(10)
            r_n.font.color.rgb = RGBColor(60, 60, 60)
        else:
            dc.paragraphs[-1].paragraph_format.space_after = Pt(6)

        # Tiny gap between items
        gap = doc.add_paragraph()
        gap.paragraph_format.space_before = Pt(0)
        gap.paragraph_format.space_after  = Pt(3)

    # ── Separator ──────────────────────────────────────────────────────────
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(16)
    sep.paragraph_format.space_after  = Pt(8)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Footer info ────────────────────────────────────────────────────────
    def _footer_row(label, value):
        if not (value or "").strip():
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r_l = p.add_run(f"{label}  ")
        r_l.bold = True
        r_l.font.size = Pt(10)
        r_l.font.color.rgb = RGBColor(60, 60, 60)
        r_v = p.add_run(str(value or ""))
        r_v.font.size = Pt(11)
        r_v.font.color.rgb = RGBColor(0, 0, 0)

    _footer_row("JOB:",  header.get("Job", ""))
    _footer_row("NOTE:", header.get("Notes", ""))

    # ── Save .docx ─────────────────────────────────────────────────────────
    doc.save(out_path)

    # ── PDF export: try reportlab first (no Word required), then Word COM ──
    pdf_path = out_path.replace(".docx", ".pdf")
    exported = _generate_bag_pdf_reportlab(header, items, pdf_path,
                                           item_start=item_start,
                                           grand_total=grand_total)

    if not exported and platform.system() == "Windows":
        exported = _export_word_pdf_via_cached_word(out_path, pdf_path)
        if not exported:
            exported = _export_word_pdf_powershell(out_path, pdf_path)

    if exported and os.path.exists(pdf_path):
        if auto_open:
            try:
                os.startfile(pdf_path)
            except Exception:
                pass
        return pdf_path

    # Fallback: open the Word doc
    if auto_open:
        if platform.system() == "Windows":
            try:
                os.startfile(out_path)
            except Exception:
                pass
        elif platform.system() == "Darwin":
            subprocess.call(["open", out_path])
        else:
            subprocess.call(["xdg-open", out_path])
    return out_path


def _generate_bag_pdf_reportlab(header: dict, items: list, pdf_path: str,
                                 item_start: int = 1,
                                 grand_total: int = None) -> bool:
    """Generate a bag docket PDF directly using reportlab. Returns True on success."""
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
    except ImportError:
        return False

    try:
        PAGE_W, PAGE_H = A4
        MARGIN  = 42.0
        C_W     = PAGE_W - 2 * MARGIN

        # Greyscale palette
        K100 = colors.HexColor("#1A1A1A")   # near-black headings
        K70  = colors.HexColor("#4D4D4D")   # sub-text
        K30  = colors.HexColor("#B3B3B3")   # rules / light text
        G10  = colors.HexColor("#E8E8E8")   # header band fill
        G05  = colors.HexColor("#F4F4F4")   # alternating row fill
        WHITE = colors.white

        c = rl_canvas.Canvas(pdf_path, pagesize=A4)

        d_ord      = header.get("Date Ordered", "") or ""
        d_due      = header.get("Date Due", "")    or ""
        cust       = (header.get("Customer Name") or "").upper()
        ord_n      = (header.get("Order Number")  or "")
        attn       = (header.get("Attention")      or "").upper()
        job        = (header.get("Job")            or "")
        loc        = (header.get("Location")       or "").upper()
        notes      = (header.get("Notes")          or "")
        created_by = (header.get("Created By")     or "").upper()
        total_shown = grand_total if grand_total is not None else len(items)

        def _font(bold=False, size=10):
            c.setFont("Helvetica-Bold" if bold else "Helvetica", size)

        def _txt(x, y, text, size=10, bold=False, col=K100, align="left"):
            _font(bold, size)
            c.setFillColor(col)
            s = str(text)
            if align == "center": c.drawCentredString(x, y, s)
            elif align == "right": c.drawRightString(x, y, s)
            else: c.drawString(x, y, s)

        def _hrule(y, lw=0.5, col=K30):
            c.setStrokeColor(col)
            c.setLineWidth(lw)
            c.line(MARGIN, y, MARGIN + C_W, y)

        def _new_page(y, need=80):
            if y < MARGIN + need:
                _draw_footer(y)
                c.showPage()
                return _draw_page_header()
            return y

        def _draw_page_header():
            y = PAGE_H - MARGIN

            # Top rule (thick)
            c.setStrokeColor(K100)
            c.setLineWidth(2)
            c.line(MARGIN, y, MARGIN + C_W, y)

            y -= 6

            # Right-side meta column — each row is "LABEL  VALUE" right-aligned
            rx  = MARGIN + C_W
            meta = []
            if d_ord:      meta.append(("DATE ORDERED", d_ord))
            if d_due:      meta.append(("DATE DUE",     d_due))
            if job:        meta.append(("JOB",          job))
            if loc:        meta.append(("LOCATION",     loc))
            if created_by: meta.append(("CREATED BY",  created_by))

            my = y - 2
            for lbl, val in meta:
                # Draw label in grey then value in black, both right-aligned as a unit
                val_w   = c.stringWidth(str(val),    "Helvetica-Bold", 8)
                gap     = 5
                lbl_str = lbl + ":"
                lbl_w   = c.stringWidth(lbl_str, "Helvetica", 7)
                total_w = lbl_w + gap + val_w
                lx = rx - total_w
                c.setFont("Helvetica", 7);      c.setFillColor(K70);  c.drawString(lx, my, lbl_str)
                c.setFont("Helvetica-Bold", 8); c.setFillColor(K100); c.drawString(lx + lbl_w + gap, my, str(val))
                my -= 11

            # Customer name — larger for visibility
            _txt(MARGIN, y - 16, cust, size=20, bold=True, col=K100)

            # O/N + ATT line below customer name — bolder and larger
            y -= 34
            parts = []
            if ord_n: parts.append(f"O/N  {ord_n}")
            if attn:  parts.append(f"ATT  {attn}")
            _txt(MARGIN, y, "   ·   ".join(parts), size=11, bold=True, col=K100)

            # Rule under header
            y -= 10
            _hrule(y, lw=0.75, col=K100)
            y -= 6

            # Column labels row
            _txt(MARGIN,      y - 10, "ITEM",        size=7, bold=True, col=K70)
            _txt(MARGIN + 46, y - 10, "DESCRIPTION", size=7, bold=True, col=K70)
            _txt(rx,          y - 10, "P/N",         size=7, bold=True, col=K70, align="right")
            _hrule(y - 14, lw=0.4, col=K30)

            return y - 18

        def _draw_footer(y):
            fy = MARGIN + 20
            _hrule(fy, lw=0.5, col=K30)
            parts = []
            if notes: parts.append(f"NOTES: {notes}")
            if parts:
                _txt(MARGIN, fy - 10, "   ".join(parts), size=8, col=K70)

        # ── First page header ────────────────────────────────────────────────
        y = _draw_page_header()

        # ── Items ────────────────────────────────────────────────────────────
        BADGE_W = 40
        PAD     = 6

        for i, item in enumerate(items):
            stock_num  = item_start + i
            label      = build_label_line(item)
            dims       = build_dims_line(item)
            pn         = (item.get("part_number") or "").strip()
            item_notes = (item.get("notes") or "").strip()

            lines = 1 + (1 if dims else 0) + (1 if item_notes else 0)
            row_h = lines * 17 + PAD * 2 + 4   # taller rows for better visibility

            y = _new_page(y, row_h + 4)

            # Row background (alternating)
            if i % 2 == 0:
                c.setFillColor(G05)
                c.rect(MARGIN, y - row_h, C_W, row_h, fill=1, stroke=0)

            # Left badge column (slightly darker)
            c.setFillColor(G10)
            c.rect(MARGIN, y - row_h, BADGE_W, row_h, fill=1, stroke=0)

            # Item number centred in badge
            _txt(MARGIN + BADGE_W / 2, y - row_h / 2 - 4,
                 str(stock_num), size=16, bold=True, col=K100, align="center")
            _txt(MARGIN + BADGE_W / 2, y - row_h + PAD,
                 f"of {total_shown}", size=7, col=K70, align="center")

            # Thin vertical rule after badge
            c.setStrokeColor(K30)
            c.setLineWidth(0.4)
            c.line(MARGIN + BADGE_W, y - row_h, MARGIN + BADGE_W, y)

            # Description block
            tx = MARGIN + BADGE_W + PAD + 2
            ty = y - PAD - 12
            _txt(tx, ty, label, size=12, bold=True, col=K100)
            if dims:
                ty -= 16
                _txt(tx, ty, dims, size=10, col=K70)
            if item_notes:
                ty -= 15
                _txt(tx, ty, f"Note: {item_notes}", size=9, col=K70)

            # P/N right-aligned
            if pn:
                pn_x = MARGIN + C_W - 4
                pn_y = y - PAD - 10
                _txt(pn_x, pn_y, pn, size=8, col=K70, align="right")
                _txt(pn_x - c.stringWidth(pn, "Helvetica", 8) - 4,
                     pn_y, "P/N", size=7, bold=True, col=K30, align="right")

            # Bottom hairline
            _hrule(y - row_h, lw=0.3, col=K30)

            y -= row_h + 1

        _draw_footer(y)
        c.save()
        return os.path.exists(pdf_path)
    except Exception:
        return False
    except Exception:
        return False


# ── Filter item display helpers ───────────────────────────────────────────────

def _build_filter_label(item: dict) -> str:
    """Label line for a panel filter item."""
    qty = item.get("Quantity", 1)
    ft  = (item.get("Filter Type") or "").strip()
    mt  = (item.get("Media Type")  or "").strip()
    label = f"{qty} × {ft}"
    if mt and mt.upper() != ft.upper():
        label += f" / {mt}"
    return label


def _build_filter_dims(item: dict) -> str:
    """Dimension string for a panel filter item (Short × Long × Channel mm)."""
    s  = int(item.get("Short",   0) or 0)
    lo = int(item.get("Long",    0) or 0)
    ch = int(item.get("Channel", 0) or 0)
    if s and lo and ch:
        return f"{s}mm × {lo}mm × {ch}mm"
    return ""


def _build_filter_options(item: dict) -> str:
    """Options flags for a panel filter item."""
    opts = []
    if item.get("Pleat Insert"):        opts.append("Pleat Insert")
    if item.get("Header"):              opts.append("Header Only")
    if item.get("Use Stock V-form"):    opts.append("Stock V-form")
    if item.get("Use Stock Flyscreen"): opts.append("Stock Flyscreen")
    return "  ·  ".join(opts)


def _build_bag_options(item: dict) -> str:
    """Options flags for a bag / roll item."""
    opts = []
    if item.get("on_wire"):      opts.append("On Wire")
    if item.get("gelled"):       opts.append("Gelled")
    if item.get("special_size"): opts.append("Special Size")
    ls = (item.get("label_suffix") or "").strip()
    if ls:                       opts.append(ls)
    return "  ·  ".join(opts)


def _set_cell_margins(cell, top: int = 0, bottom: int = 0,
                      left: int = 0, right: int = 0):
    """Set table-cell internal margins in DXA (twentieths of a point)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _set_cell_valign(cell, val: str = "center"):
    """Set vertical alignment of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


# ── Unified order docket ──────────────────────────────────────────────────────

def generate_unified_docket(header: dict, items: List[dict], out_path: str) -> str:
    """
    Build a single Word docket (.docx) covering the entire order —
    panel filters, bag filters, and media rolls together.
    Every item is stamped with its stock number (X / N).
    Saves to out_path, attempts PDF export via PowerShell Word COM,
    opens the result, and returns the opened path.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed.  Run:  pip install python-docx")

    doc   = Document()
    total = len(items)

    # ── Page setup (A4 portrait) ───────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    content_w = int(section.page_width - section.left_margin - section.right_margin)
    navy_hex  = "1A4F8A"

    # ── Header banner table ────────────────────────────────────────────────
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.style   = "Table Grid"
    hdr_tbl.autofit = False
    left_w  = int(content_w * 0.65)
    right_w = content_w - left_w
    hdr_tbl.columns[0].width = left_w
    hdr_tbl.columns[1].width = right_w

    lc = hdr_tbl.rows[0].cells[0]
    rc = hdr_tbl.rows[0].cells[1]
    for c in (lc, rc):
        _set_cell_bg(c, navy_hex)
        _remove_cell_borders(c)

    # Left: customer name + O/N + ATT + location
    lc_p0 = lc.paragraphs[0]
    lc_p0.paragraph_format.space_before = Pt(6)
    lc_p0.paragraph_format.space_after  = Pt(2)
    r = lc_p0.add_run((header.get("Customer Name") or "").upper())
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(255, 255, 255)

    def _lc_row(lbl, val):
        p = lc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        rl = p.add_run(f"{lbl}  ")
        rl.bold = True; rl.font.size = Pt(10)
        rl.font.color.rgb = RGBColor(169, 204, 227)
        rv = p.add_run(str(val or ""))
        rv.font.size = Pt(10)
        rv.font.color.rgb = RGBColor(255, 255, 255)

    _lc_row("O/N:", header.get("Order Number", ""))
    _lc_row("ATT:", header.get("Attention", ""))
    if header.get("Location"):
        _lc_row("LOCATION:", header.get("Location", ""))
    lc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Right: dates + job + total item count
    _rc_first = [True]

    def _rc_row(lbl, val):
        p = rc.paragraphs[0] if _rc_first[0] else rc.add_paragraph()
        if _rc_first[0]:
            _rc_first[0] = False
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rl = p.add_run(f"{lbl}  ")
        rl.bold = True; rl.font.size = Pt(10)
        rl.font.color.rgb = RGBColor(169, 204, 227)
        rv = p.add_run(str(val or ""))
        rv.font.size = Pt(10)
        rv.font.color.rgb = RGBColor(255, 255, 255)

    _rc_row("DATE:", header.get("Date Ordered", ""))
    _rc_row("DUE:",  header.get("Date Due",     ""))
    if header.get("Job"):
        _rc_row("JOB:", header.get("Job", ""))

    p_tot = rc.add_paragraph()
    p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_tot.paragraph_format.space_before = Pt(4)
    p_tot.paragraph_format.space_after  = Pt(6)
    rt = p_tot.add_run(f"TOTAL  {total} ITEM{'S' if total != 1 else ''}")
    rt.bold = True; rt.font.size = Pt(11)
    rt.font.color.rgb = RGBColor(255, 255, 255)

    # Spacer after header
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(6)
    sp.paragraph_format.space_after  = Pt(4)

    # ── Item blocks ────────────────────────────────────────────────────────
    STOCK_W  = int(Cm(1.6))           # stock-number column (EMU)
    DETAIL_W = content_w - STOCK_W    # detail column       (EMU)
    _BG_EVEN = "EBF5FB"               # very light blue
    _BG_ODD  = "FFFFFF"               # white

    for idx, item in enumerate(items):
        stock_num = idx + 1
        kind      = item.get("item_kind", "filter")

        if kind == "bag":
            label = build_label_line(item)
            dims  = build_dims_line(item)
            pn    = (item.get("part_number") or "").strip()
            opts  = _build_bag_options(item)
            notes = (item.get("notes") or "").strip()
        else:
            label = _build_filter_label(item)
            dims  = _build_filter_dims(item)
            pn    = ""
            opts  = _build_filter_options(item)
            notes = (item.get("Notes") or "").strip()

        # Build the per-item 2-column table
        itbl = doc.add_table(rows=1, cols=2)
        itbl.style   = "Table Grid"
        itbl.autofit = False
        itbl.columns[0].width = STOCK_W
        itbl.columns[1].width = DETAIL_W

        sc = itbl.rows[0].cells[0]   # left: stock number
        dc = itbl.rows[0].cells[1]   # right: item details

        # ── Stock-number cell (navy) ──────────────────────────────────────
        _set_cell_bg(sc, navy_hex)
        _remove_cell_borders(sc)
        _set_cell_valign(sc, "center")
        _set_cell_margins(sc, top=80, bottom=80, left=60, right=60)

        sc_p1 = sc.paragraphs[0]
        sc_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc_p1.paragraph_format.space_before = Pt(0)
        sc_p1.paragraph_format.space_after  = Pt(0)
        r_cur = sc_p1.add_run(str(stock_num))
        r_cur.bold = True; r_cur.font.size = Pt(18)
        r_cur.font.color.rgb = RGBColor(255, 255, 255)

        sc_p2 = sc.add_paragraph()
        sc_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc_p2.paragraph_format.space_before = Pt(0)
        sc_p2.paragraph_format.space_after  = Pt(0)
        r_sl = sc_p2.add_run("/")
        r_sl.font.size = Pt(9)
        r_sl.font.color.rgb = RGBColor(169, 204, 227)

        sc_p3 = sc.add_paragraph()
        sc_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc_p3.paragraph_format.space_before = Pt(0)
        sc_p3.paragraph_format.space_after  = Pt(0)
        r_tot = sc_p3.add_run(str(total))
        r_tot.bold = True; r_tot.font.size = Pt(12)
        r_tot.font.color.rgb = RGBColor(200, 220, 240)

        # ── Detail cell ───────────────────────────────────────────────────
        bg = _BG_EVEN if (idx % 2 == 0) else _BG_ODD
        _set_cell_bg(dc, bg)
        _remove_cell_borders(dc)
        _set_cell_margins(dc, top=100, bottom=100, left=180, right=120)

        # Label (bold, navy, 12 pt)
        dc_p0 = dc.paragraphs[0]
        dc_p0.paragraph_format.space_before = Pt(2)
        dc_p0.paragraph_format.space_after  = Pt(2)
        rl = dc_p0.add_run(label)
        rl.bold = True; rl.font.size = Pt(12)
        rl.font.color.rgb = RGBColor(*_NAVY)

        # Dimensions
        if dims:
            dp = dc.add_paragraph()
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after  = Pt(2)
            rd = dp.add_run(dims)
            rd.font.size = Pt(11)
            rd.font.color.rgb = RGBColor(45, 55, 72)

        # P/N (bags only)
        if pn:
            pp = dc.add_paragraph()
            pp.paragraph_format.space_before = Pt(0)
            pp.paragraph_format.space_after  = Pt(2)
            r_pl = pp.add_run("P/N:  ")
            r_pl.bold = True; r_pl.font.size = Pt(10)
            r_pl.font.color.rgb = RGBColor(113, 128, 150)
            r_pv = pp.add_run(pn)
            r_pv.font.size = Pt(10)
            r_pv.font.color.rgb = RGBColor(45, 55, 72)

        # Options
        if opts:
            op = dc.add_paragraph()
            op.paragraph_format.space_before = Pt(0)
            op.paragraph_format.space_after  = Pt(2)
            r_ol = op.add_run("Options:  ")
            r_ol.bold = True; r_ol.font.size = Pt(10)
            r_ol.font.color.rgb = RGBColor(113, 128, 150)
            r_ov = op.add_run(opts)
            r_ov.font.size = Pt(10)
            r_ov.font.color.rgb = RGBColor(45, 55, 72)

        # Notes
        if notes:
            np_ = dc.add_paragraph()
            np_.paragraph_format.space_before = Pt(0)
            np_.paragraph_format.space_after  = Pt(4)
            r_n = np_.add_run(f"Note:  {notes}")
            r_n.italic = True; r_n.font.size = Pt(10)
            r_n.font.color.rgb = RGBColor(113, 128, 150)
        else:
            dc.paragraphs[-1].paragraph_format.space_after = Pt(6)

        # Tiny gap between item blocks
        gap = doc.add_paragraph()
        gap.paragraph_format.space_before = Pt(0)
        gap.paragraph_format.space_after  = Pt(3)

    # ── Divider ────────────────────────────────────────────────────────────
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(10)
    sep.paragraph_format.space_after  = Pt(6)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A4F8A")
    pBdr.append(bot)
    pPr.append(pBdr)

    # ── Footer (job / order notes) ─────────────────────────────────────────
    def _foot(lbl, val):
        if not str(val or "").strip():
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r_l = p.add_run(f"{lbl}  ")
        r_l.bold = True; r_l.font.size = Pt(10)
        r_l.font.color.rgb = RGBColor(113, 128, 150)
        r_v = p.add_run(str(val))
        r_v.font.size = Pt(11)
        r_v.font.color.rgb = RGBColor(45, 55, 72)

    _foot("JOB:",  header.get("Job",   ""))
    _foot("NOTE:", header.get("Notes", ""))

    # ── Save .docx ─────────────────────────────────────────────────────────
    doc.save(out_path)

    # ── Export PDF via PowerShell Word COM ─────────────────────────────────
    if platform.system() == "Windows":
        pdf_path = out_path.replace(".docx", ".pdf")
        if _export_word_pdf_powershell(out_path, pdf_path) and os.path.exists(pdf_path):
            os.startfile(pdf_path)
            return pdf_path

    # Fallback: open the Word doc
    if platform.system() == "Windows":
        os.startfile(out_path)
    elif platform.system() == "Darwin":
        subprocess.call(["open", out_path])
    else:
        subprocess.call(["xdg-open", out_path])
    return out_path


def _export_word_pdf_powershell(docx_path: str, pdf_path: str) -> bool:
    """Export a .docx to .pdf using PowerShell Word COM. Returns True on success."""
    docx_ps = docx_path.replace("'", "''")
    pdf_ps  = pdf_path.replace("'", "''")
    script = (
        "$ErrorActionPreference='Stop';"
        "$w=New-Object -ComObject Word.Application;"
        "$w.Visible=$false;"
        f"$d=$w.Documents.Open('{docx_ps}');"
        f"$d.SaveAs2('{pdf_ps}',17);"   # 17 = wdFormatPDF
        "$d.Close($false);"
        "$w.Quit();"
        "[System.Runtime.InteropServices.Marshal]::ReleaseComObject($w)|Out-Null"
    )
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
            capture_output=True, timeout=90,
            creationflags=subprocess.CREATE_NO_WINDOW,
        )
        return result.returncode == 0 and os.path.exists(pdf_path)
    except Exception:
        return False
